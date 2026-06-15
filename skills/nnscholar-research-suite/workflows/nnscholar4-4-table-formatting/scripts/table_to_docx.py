#!/usr/bin/env python
"""Create Word three-line manuscript tables from structured table files."""

from __future__ import annotations

import argparse
import html
import json
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def read_table(path: Path, sheet: str | None = None) -> pd.DataFrame:
    suffix = path.suffix.lower()
    if suffix == ".csv":
        return pd.read_csv(path)
    if suffix in {".tsv", ".txt"}:
        return pd.read_csv(path, sep="\t")
    if suffix in {".xlsx", ".xls"}:
        return pd.read_excel(path, sheet_name=sheet or 0)
    if suffix == ".json":
        return pd.read_json(path)
    if suffix in {".md", ".markdown"}:
        return read_markdown_table(path)
    raise ValueError(f"Unsupported input file type: {suffix}")


def read_markdown_table(path: Path) -> pd.DataFrame:
    lines = [line.strip() for line in path.read_text(encoding="utf-8-sig").splitlines() if line.strip()]
    table_lines = [line for line in lines if line.startswith("|") and line.endswith("|")]
    if len(table_lines) < 2:
        raise ValueError("Markdown input must contain a pipe table.")
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if all(set(cell.replace(":", "").replace("-", "")) == set() for cell in cells):
            continue
        if all(cell.replace(":", "").replace("-", "").strip() == "" for cell in cells):
            continue
        rows.append(cells)
    if len(rows) < 2:
        raise ValueError("Markdown table must contain header and at least one row.")
    header = rows[0]
    body = [row + [""] * (len(header) - len(row)) for row in rows[1:]]
    return pd.DataFrame(body, columns=header)


def clean_df(df: pd.DataFrame, max_rows: int | None = None) -> pd.DataFrame:
    df = df.copy()
    df.columns = [str(col).strip() for col in df.columns]
    df = df.fillna("")
    if max_rows is not None:
        df = df.head(max_rows)
    return df


def set_cell_text(cell: Any, text: str, font: str, size_pt: float, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_cell_borders(cell: Any, **kwargs: dict[str, str]) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        if edge_data is None:
            element.set(qn("w:val"), "nil")
        else:
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))


def apply_three_line_borders(table: Any) -> None:
    top_border = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
    mid_border = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    bottom_border = {"val": "single", "sz": "12", "space": "0", "color": "000000"}

    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell)

    for cell in table.rows[0].cells:
        set_cell_borders(cell, top=top_border, bottom=mid_border)
    for cell in table.rows[-1].cells:
        set_cell_borders(cell, bottom=bottom_border)


def write_docx(df: pd.DataFrame, output: Path, args: argparse.Namespace) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    style = doc.styles["Normal"]
    style.font.name = args.font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), args.font)
    style.font.size = Pt(args.font_size)

    if args.caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(args.caption)
        run.bold = True
        run.font.name = args.font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), args.font)
        run.font.size = Pt(args.font_size)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.autofit = True
    for idx, col in enumerate(df.columns):
        set_cell_text(table.rows[0].cells[idx], col, args.font, args.font_size, bold=True)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(df.columns):
            set_cell_text(cells[idx], row[col], args.font, args.font_size, bold=False)

    if args.style == "three-line":
        apply_three_line_borders(table)

    notes = build_notes(args)
    for note in notes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(note)
        run.font.name = args.font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), args.font)
        run.font.size = Pt(max(args.font_size - 1, 8))

    doc.save(output)


def build_notes(args: argparse.Namespace) -> list[str]:
    notes: list[str] = []
    if args.note:
        notes.extend(args.note)
    if args.abbreviations:
        notes.append(f"Abbreviations: {args.abbreviations}")
    if args.statistics:
        notes.append(f"Statistical notes: {args.statistics}")
    return notes


def to_markdown(df: pd.DataFrame, caption: str | None, notes: list[str]) -> str:
    lines: list[str] = []
    if caption:
        lines.append(f"**{caption}**")
        lines.append("")
    lines.append("| " + " | ".join(str(col) for col in df.columns) + " |")
    lines.append("| " + " | ".join("---" for _ in df.columns) + " |")
    for _, row in df.iterrows():
        lines.append("| " + " | ".join(str(row[col]) for col in df.columns) + " |")
    if notes:
        lines.append("")
        lines.extend(notes)
    return "\n".join(lines) + "\n"


def to_html(df: pd.DataFrame, caption: str | None, notes: list[str]) -> str:
    caption_html = f"<caption>{html.escape(caption)}</caption>" if caption else ""
    header = "".join(f"<th>{html.escape(str(col))}</th>" for col in df.columns)
    rows = []
    for _, row in df.iterrows():
        rows.append("<tr>" + "".join(f"<td>{html.escape(str(row[col]))}</td>" for col in df.columns) + "</tr>")
    notes_html = "".join(f"<p class=\"table-note\">{html.escape(note)}</p>" for note in notes)
    return f"""<!doctype html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{ font-family: Arial, sans-serif; }}
table.three-line {{ border-collapse: collapse; width: 100%; font-size: 11pt; }}
table.three-line caption {{ caption-side: top; text-align: left; font-weight: bold; margin-bottom: 8px; }}
table.three-line thead tr {{ border-top: 2px solid #000; border-bottom: 1px solid #000; }}
table.three-line tbody tr:last-child {{ border-bottom: 2px solid #000; }}
table.three-line th, table.three-line td {{ padding: 6px 8px; text-align: left; }}
.table-note {{ font-size: 10pt; margin: 6px 0; }}
</style>
</head>
<body>
<table class="three-line">
{caption_html}
<thead><tr>{header}</tr></thead>
<tbody>
{''.join(rows)}
</tbody>
</table>
{notes_html}
</body>
</html>
"""


def write_manifest(output_dir: Path, args: argparse.Namespace, outputs: list[str], df: pd.DataFrame) -> None:
    manifest = {
        "input": str(args.input),
        "outputs": outputs,
        "caption": args.caption,
        "notes": build_notes(args),
        "shape": {"rows": int(df.shape[0]), "columns": int(df.shape[1])},
        "columns": list(df.columns),
    }
    (output_dir / "table_manifest.json").write_text(json.dumps(manifest, indent=2, ensure_ascii=False), encoding="utf-8")


def write_report(output_dir: Path, args: argparse.Namespace, outputs: list[str], df: pd.DataFrame) -> Path:
    notes = build_notes(args)
    lines = [
        "# Manuscript Table Report",
        "",
        "## Table Title",
        "",
        args.caption or "Table X. [Add a concise title that states the table purpose.]",
        "",
        "## Table Structure",
        "",
        f"- Rows: {df.shape[0]}",
        f"- Columns: {df.shape[1]}",
        f"- Column headers: {', '.join(str(col) for col in df.columns)}",
        "",
        "## Table Notes",
        "",
    ]
    if notes:
        lines.extend(f"- {note}" for note in notes)
    else:
        lines.append("- Note. [Add summary format, denominator, missingness, or source note when relevant.]")
        lines.append("- Abbreviations: [Define all nonstandard abbreviations.]")
        lines.append("- Statistical notes: [State tests, models, uncertainty intervals, and corrections when relevant.]")
    lines.extend(
        [
            "",
            "## Quality Audit",
            "",
            "- Header clarity: check that each column name is self-explanatory and units are visible.",
            "- Units and denominators: add n, %, time windows, scale ranges, or denominators where relevant.",
            "- Decimal consistency: align precision across comparable values.",
            "- Abbreviations: define every nonstandard abbreviation in the notes.",
            "- Statistical explanation: state test/model, adjustment, uncertainty interval, and p-value policy.",
            "- Claim support: make sure the title and notes do not overstate what the table proves.",
            "",
            "## Generated Files",
            "",
        ]
    )
    lines.extend(f"- {output}" for output in outputs)
    report_path = output_dir / "table_report.md"
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return report_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Create manuscript table outputs, including Word three-line tables.")
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--sheet", default=None)
    parser.add_argument("--output-dir", required=True, type=Path)
    parser.add_argument("--basename", default="table")
    parser.add_argument("--caption", default=None)
    parser.add_argument("--note", action="append", default=[])
    parser.add_argument("--abbreviations", default=None)
    parser.add_argument("--statistics", default=None)
    parser.add_argument("--style", choices=["three-line", "plain"], default="three-line")
    parser.add_argument("--formats", default="docx,md,html")
    parser.add_argument("--font", default="Arial")
    parser.add_argument("--font-size", type=float, default=10.5)
    parser.add_argument("--max-rows", type=int, default=None)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    df = clean_df(read_table(args.input, args.sheet), args.max_rows)
    args.output_dir.mkdir(parents=True, exist_ok=True)
    formats = {item.strip().lower() for item in args.formats.split(",") if item.strip()}
    notes = build_notes(args)
    outputs: list[str] = []

    if "docx" in formats:
        output = args.output_dir / f"{args.basename}.docx"
        write_docx(df, output, args)
        outputs.append(str(output))
    if "md" in formats or "markdown" in formats:
        output = args.output_dir / f"{args.basename}.md"
        output.write_text(to_markdown(df, args.caption, notes), encoding="utf-8")
        outputs.append(str(output))
    if "html" in formats:
        output = args.output_dir / f"{args.basename}.html"
        output.write_text(to_html(df, args.caption, notes), encoding="utf-8")
        outputs.append(str(output))

    write_manifest(args.output_dir, args, outputs, df)
    report = write_report(args.output_dir, args, outputs, df)
    print(json.dumps({"outputs": outputs, "manifest": str(args.output_dir / "table_manifest.json"), "report": str(report)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
